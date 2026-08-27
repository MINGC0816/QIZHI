from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".markdown", ".txt"}


def load_file(path: Path) -> list[Document]:
    """按后缀解析单个本地文件为 Document 列表。"""
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(path)

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"不支持的文件类型: {suffix}，支持 {sorted(SUPPORTED_SUFFIXES)}")

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    return _load_text(path)


def load_paths(paths: list[Path]) -> list[Document]:
    docs: list[Document] = []
    for p in paths:
        docs.extend(load_file(p))
    return docs


def _load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    docs: list[Document] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "source_path": str(path.resolve()),
                    "page": i,
                    "file_type": "pdf",
                },
            )
        )
    if not docs:
        raise ValueError(f"PDF 无可提取文本（可能是扫描件）: {path}")
    return docs


def _load_docx(path: Path) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ValueError(f"Word 文档无有效文本: {path}")
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "source_path": str(path.resolve()),
                "page": 1,
                "file_type": "docx",
            },
        )
    ]


def _load_text(path: Path) -> list[Document]:
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        raise ValueError(f"空文件: {path}")
    suffix = path.suffix.lower().lstrip(".") or "txt"
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "source_path": str(path.resolve()),
                "page": 1,
                "file_type": suffix,
            },
        )
    ]
